"""Re-run the PoC proposal edit and display a clean before/after diff.

Outputs are saved to ./output/ so you can open them in Word.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from poc_proposal_test import build_poc_template
from report_mcp.server import fill_and_save, list_template_targets

PROJECT = Path(__file__).parent
OUT_DIR = PROJECT / "output"


def find(targets: list[dict], needle: str, kind: str) -> dict:
    for t in targets:
        if t.get("target_kind") == kind and needle in (t.get("current_text") or ""):
            return t
    raise SystemExit(f"target not found: {needle!r} ({kind})")


def dump_doc(label: str, path: Path) -> None:
    print(f"\n=== {label}  ({path}) ===")
    d = Document(path)
    for i, p in enumerate(d.paragraphs):
        if p.text.strip():
            print(f"  ¶{i:02d} | {p.text}")
    for ti, table in enumerate(d.tables):
        print(f"  ┌─ table[{ti}] ─────────────")
        for ri, row in enumerate(table.rows):
            cells = [c.text for c in row.cells]
            print("  │  " + " | ".join(f"{c:<22}" for c in cells))
        print("  └──────────────────────────")


def main() -> int:
    OUT_DIR.mkdir(exist_ok=True)
    template = OUT_DIR / "poc_template.docx"
    filled = OUT_DIR / "poc_filled.docx"

    print(f"[1] building original PoC proposal template -> {template}")
    build_poc_template(template)

    print(f"[2] discovering editable targets")
    paras = list_template_targets(str(template), target_kinds=["paragraph"], max_targets=200)["targets"]
    cells = list_template_targets(str(template), target_kinds=["cell"], max_targets=200)["targets"]

    cover    = find(paras, "PoC 제안서",        "paragraph")
    overview = find(paras, "[PROJECT_GOAL]",   "paragraph")
    benefit1 = find(paras, "[BENEFIT_1]",      "paragraph")
    benefit2 = find(paras, "[BENEFIT_2]",      "paragraph")
    benefit3 = find(paras, "[BENEFIT_3]",      "paragraph")
    subtitle = find(paras, "[SUBMIT_DATE]",    "paragraph")
    item1    = find(cells, "[ITEM_1]",  "cell")
    qty1     = find(cells, "[QTY_1]",   "cell")
    price1   = find(cells, "[PRICE_1]", "cell")
    item2    = find(cells, "[ITEM_2]",  "cell")
    qty2     = find(cells, "[QTY_2]",   "cell")
    price2   = find(cells, "[PRICE_2]", "cell")
    total    = find(cells, "[TOTAL]",   "cell")

    edits = [
        # Cover + subtitle
        {"edit_type": "text", "target_kind": "paragraph",
         "target_id": cover["target_id"], "expected_text_hash": cover["text_hash"],
         "new_text": "사내 AI 챗봇 도입 PoC 제안서"},
        {"edit_type": "text", "target_kind": "paragraph",
         "target_id": subtitle["target_id"], "expected_text_hash": subtitle["text_hash"],
         "new_text": "제출일: 2026-05-15 · 작성: iHopper AI팀"},
        # Overview body
        {"edit_type": "text", "target_kind": "paragraph",
         "target_id": overview["target_id"], "expected_text_hash": overview["text_hash"],
         "new_text": ("본 PoC는 사내 문의 응대 자동화 효과를 정량 검증하는 것을 목적으로 하며, "
                      "8주 동안 영업·CS팀 50명을 대상으로 수행됩니다.")},
        # Pricing table (every data cell)
        {"edit_type": "text", "target_kind": "cell",
         "target_id": item1["target_id"], "expected_text_hash": item1["text_hash"],
         "new_text": "챗봇 모델 API 사용료"},
        {"edit_type": "text", "target_kind": "cell",
         "target_id": qty1["target_id"], "expected_text_hash": qty1["text_hash"],
         "new_text": "8주"},
        {"edit_type": "text", "target_kind": "cell",
         "target_id": price1["target_id"], "expected_text_hash": price1["text_hash"],
         "new_text": "12,000,000"},
        {"edit_type": "text", "target_kind": "cell",
         "target_id": item2["target_id"], "expected_text_hash": item2["text_hash"],
         "new_text": "구축·운영 인건비"},
        {"edit_type": "text", "target_kind": "cell",
         "target_id": qty2["target_id"], "expected_text_hash": qty2["text_hash"],
         "new_text": "1식"},
        {"edit_type": "text", "target_kind": "cell",
         "target_id": price2["target_id"], "expected_text_hash": price2["text_hash"],
         "new_text": "28,000,000"},
        {"edit_type": "text", "target_kind": "cell",
         "target_id": total["target_id"], "expected_text_hash": total["text_hash"],
         "new_text": "40,000,000"},
        # Benefits
        {"edit_type": "text", "target_kind": "paragraph",
         "target_id": benefit1["target_id"], "expected_text_hash": benefit1["text_hash"],
         "new_text": "· 1차 문의 응대 평균 처리 시간 6분 → 30초 미만"},
        {"edit_type": "text", "target_kind": "paragraph",
         "target_id": benefit2["target_id"], "expected_text_hash": benefit2["text_hash"],
         "new_text": "· 야간·휴일 응대 공백 해소 (24/7 운영)"},
        {"edit_type": "text", "target_kind": "paragraph",
         "target_id": benefit3["target_id"], "expected_text_hash": benefit3["text_hash"],
         "new_text": "· FAQ 자동 학습으로 신규 직원 온보딩 시간 40% 단축"},
    ]

    print(f"[3] applying {len(edits)} edits and saving -> {filled}")
    result = fill_and_save(str(template), edits, str(filled))
    if result.get("status") != "ok":
        print(f"FAIL: {result}")
        return 1

    dump_doc("BEFORE  (original template)", template)
    dump_doc("AFTER   (filled output)",     filled)

    # Style preservation spot-check
    pre = Document(template).paragraphs[0]
    post = Document(filled).paragraphs[0]
    pre_r, post_r = pre.runs[0], post.runs[0]
    print()
    print("=== style preservation on cover paragraph ===")
    print(f"  alignment : before={pre.alignment}   after={post.alignment}")
    print(f"  bold      : before={pre_r.bold}            after={post_r.bold}")
    print(f"  size      : before={pre_r.font.size}   after={post_r.font.size}")
    print(f"  color     : before={pre_r.font.color.rgb}  after={post_r.font.color.rgb}")

    print()
    print(f"파일 경로:")
    print(f"  원본:  {template}")
    print(f"  결과:  {filled}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
