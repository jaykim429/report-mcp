"""Realistic test: PoC-proposal-style DOCX template.

Builds a multi-section proposal (cover heading, sections, a pricing table)
with various styles, then drives it through report-mcp:
  - replace the cover title text
  - replace a section's body paragraph
  - replace a table cell's text
Then re-opens the output and asserts both content changes AND structural /
stylistic preservation (table dimensions, alignment, bold, font color).
"""

from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from report_mcp.server import (
    fill_and_save,
    list_template_targets,
)


def build_poc_template(path: Path) -> None:
    doc = Document()

    # ── Cover ─────────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cover.add_run("[PROJECT_NAME] PoC 제안서")
    cr.bold = True
    cr.font.size = Pt(24)
    cr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("제출일: [SUBMIT_DATE] · 작성: [AUTHOR_ORG]")
    sr.italic = True
    sr.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # ── Section 1 ────────────────────────────────────────────────────────
    doc.add_heading("1. 프로젝트 개요", level=1)
    p_overview = doc.add_paragraph(
        "본 PoC는 [PROJECT_GOAL]을 검증하는 것을 목적으로 하며, "
        "[DURATION] 동안 [SCOPE] 범위에서 수행됩니다."
    )
    p_overview.paragraph_format.first_line_indent = Pt(12)

    # ── Section 2: pricing table ─────────────────────────────────────────
    doc.add_heading("2. 견적 요약", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Light Grid Accent 1"
    # header
    hdr = table.rows[0].cells
    for i, label in enumerate(["항목", "수량", "금액(원)"]):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(label)
        run.bold = True
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    rows_data = [
        ("[ITEM_1]", "[QTY_1]", "[PRICE_1]"),
        ("[ITEM_2]", "[QTY_2]", "[PRICE_2]"),
        ("합계", "", "[TOTAL]"),
    ]
    for row_idx, row_data in enumerate(rows_data, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    # ── Section 3 ────────────────────────────────────────────────────────
    doc.add_heading("3. 기대 효과", level=1)
    doc.add_paragraph("· [BENEFIT_1]")
    doc.add_paragraph("· [BENEFIT_2]")
    doc.add_paragraph("· [BENEFIT_3]")

    doc.save(path)


def find_target_by_text(targets: list[dict], substring: str, kind: str) -> dict | None:
    for t in targets:
        if t.get("target_kind") == kind and substring in (t.get("current_text") or ""):
            return t
    return None


def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="poc-test-"))
    template = tmp / "poc_template.docx"
    output = tmp / "poc_filled.docx"

    print(f"[setup] building PoC proposal template -> {template}")
    build_poc_template(template)

    # Snapshot pre-edit visible state we care about
    pre = Document(template)
    pre_tables = len(pre.tables)
    pre_table_dims = (len(pre.tables[0].rows), len(pre.tables[0].columns))
    pre_cover_run = pre.paragraphs[0].runs[0]
    pre_cover_style = {
        "bold": pre_cover_run.bold,
        "size": pre_cover_run.font.size,
        "color": str(pre_cover_run.font.color.rgb) if pre_cover_run.font.color and pre_cover_run.font.color.rgb else None,
        "alignment": pre.paragraphs[0].alignment,
    }
    print(f"[setup] pre tables={pre_tables} dims={pre_table_dims} cover_style={pre_cover_style}")

    # Discover targets across paragraphs AND cells
    para_targets = list_template_targets(str(template), target_kinds=["paragraph"], max_targets=200)
    cell_targets = list_template_targets(str(template), target_kinds=["cell"], max_targets=200)

    cover_tgt = find_target_by_text(para_targets.get("targets", []), "PoC 제안서", "paragraph")
    overview_tgt = find_target_by_text(para_targets.get("targets", []), "[PROJECT_GOAL]", "paragraph")
    item1_cell = find_target_by_text(cell_targets.get("targets", []), "[ITEM_1]", "cell")
    price1_cell = find_target_by_text(cell_targets.get("targets", []), "[PRICE_1]", "cell")

    for name, tgt in [("cover", cover_tgt), ("overview", overview_tgt), ("item1_cell", item1_cell), ("price1_cell", price1_cell)]:
        if tgt is None:
            print(f"FAIL: could not locate target '{name}'")
            print(f"para targets count: {len(para_targets.get('targets', []))}")
            print(f"cell targets count: {len(cell_targets.get('targets', []))}")
            return 1
        print(f"  found {name}: id={tgt['target_id']} text={tgt.get('current_text')!r}")

    edits = [
        {
            "edit_type": "text",
            "target_kind": "paragraph",
            "target_id": cover_tgt["target_id"],
            "expected_text_hash": cover_tgt["text_hash"],
            "new_text": "AI 챗봇 도입 PoC 제안서",
        },
        {
            "edit_type": "text",
            "target_kind": "paragraph",
            "target_id": overview_tgt["target_id"],
            "expected_text_hash": overview_tgt["text_hash"],
            "new_text": (
                "본 PoC는 사내 문의 응대 자동화의 효과를 검증하는 것을 목적으로 하며, "
                "8주 동안 영업·CS팀 50명 범위에서 수행됩니다."
            ),
        },
        {
            "edit_type": "text",
            "target_kind": "cell",
            "target_id": item1_cell["target_id"],
            "expected_text_hash": item1_cell["text_hash"],
            "new_text": "챗봇 모델 API 사용료",
        },
        {
            "edit_type": "text",
            "target_kind": "cell",
            "target_id": price1_cell["target_id"],
            "expected_text_hash": price1_cell["text_hash"],
            "new_text": "12,000,000",
        },
    ]

    print(f"[apply] fill_and_save -> {output}")
    res = fill_and_save(
        template_path=str(template),
        edits=edits,
        output_path=str(output),
        dry_run=False,
    )
    print(json.dumps({k: v for k, v in res.items() if k != "result"}, ensure_ascii=False, indent=2))
    if res.get("status") != "ok":
        print(f"FAIL: fill_and_save result -> {res}")
        return 1
    if not output.is_file():
        print("FAIL: output file missing")
        return 1

    # Verify post-edit state
    post = Document(output)
    body_text = "\n".join(p.text for p in post.paragraphs)
    print("--- output paragraphs ---")
    print(body_text)
    print("--- output table[0] ---")
    for r in post.tables[0].rows:
        print(" | ".join(c.text for c in r.cells))
    print("--- end ---")

    failures: list[str] = []

    # Content checks
    if "AI 챗봇 도입 PoC 제안서" not in body_text:
        failures.append("cover title was not replaced")
    if "8주 동안 영업·CS팀 50명" not in body_text:
        failures.append("overview body was not replaced")

    table_texts = [c.text for r in post.tables[0].rows for c in r.cells]
    if "챗봇 모델 API 사용료" not in table_texts:
        failures.append("item cell was not replaced")
    if "12,000,000" not in table_texts:
        failures.append("price cell was not replaced")

    # Structural preservation
    if len(post.tables) != pre_tables:
        failures.append(f"table count changed: {pre_tables} -> {len(post.tables)}")
    post_dims = (len(post.tables[0].rows), len(post.tables[0].columns))
    if post_dims != pre_table_dims:
        failures.append(f"table dims changed: {pre_table_dims} -> {post_dims}")

    # Unchanged cells still intact
    if "[ITEM_2]" not in table_texts:
        failures.append("unedited cell [ITEM_2] was lost")
    if "합계" not in table_texts:
        failures.append("unedited cell '합계' was lost")

    # Cover style preserved despite text replacement
    post_cover_run = post.paragraphs[0].runs[0]
    post_cover_style = {
        "bold": post_cover_run.bold,
        "size": post_cover_run.font.size,
        "color": str(post_cover_run.font.color.rgb) if post_cover_run.font.color and post_cover_run.font.color.rgb else None,
        "alignment": post.paragraphs[0].alignment,
    }
    if pre_cover_style != post_cover_style:
        failures.append(f"cover style changed: {pre_cover_style} -> {post_cover_style}")
    else:
        print(f"[ok] cover style preserved exactly: {post_cover_style}")

    if failures:
        print()
        print("FAILED CHECKS:")
        for f in failures:
            print(f"  - {f}")
        return 1

    print()
    print("PASS: PoC-proposal template edited successfully —")
    print("      cover title, body paragraph, AND table cells were replaced")
    print("      while table dimensions, untouched cells, and cover styles")
    print("      (bold / 24pt / #1F4E79 / CENTER) all stayed intact.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
