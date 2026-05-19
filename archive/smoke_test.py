"""End-to-end smoke test for report-mcp.

1. Build a sample DOCX template with python-docx.
2. Call inspect_template to discover a target paragraph + its hash.
3. Call fill_and_save with a TextEdit replacing that paragraph's body.
4. Re-open the output DOCX and verify the new text is present.
"""

from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from report_mcp.server import (
    fill_and_save,
    inspect_template,
    list_template_targets,
)


def build_sample_template(path: Path) -> None:
    doc = Document()
    doc.add_heading("월간 보고서", level=0)
    doc.add_paragraph("작성자: {{author}}")
    doc.add_paragraph("요약: 이번 달의 주요 이슈는 아직 작성되지 않았습니다.")
    doc.add_heading("세부 내용", level=1)
    doc.add_paragraph("자세한 내용을 여기에 적어주세요.")
    doc.save(path)


def first_text_target(targets_result: dict) -> dict | None:
    for tgt in targets_result.get("targets", []):
        if tgt.get("target_kind") == "paragraph" and (tgt.get("current_text") or "").startswith("요약:"):
            return tgt
    return None


def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="report-mcp-smoke-"))
    template_path = tmp / "template.docx"
    output_path = tmp / "filled.docx"

    print(f"[1/5] sample template -> {template_path}")
    build_sample_template(template_path)

    print(f"[2/5] inspect_template")
    inspected = inspect_template(str(template_path), start=0, limit=20)
    print(json.dumps(inspected, ensure_ascii=False, indent=2)[:1200])

    print(f"[3/5] list_template_targets")
    targets = list_template_targets(str(template_path), target_kinds=["paragraph"], max_targets=50)
    print(json.dumps(targets, ensure_ascii=False, indent=2)[:1200])

    tgt = first_text_target(targets)
    if tgt is None:
        print("FAIL: no editable paragraph target returned")
        return 1
    target_id = tgt["target_id"]
    expected_hash = tgt.get("text_hash")
    if not expected_hash:
        print(f"FAIL: could not locate text_hash for target {target_id}")
        print(f"target keys: {list(tgt.keys())}")
        return 1

    new_text = "요약: 매출 12% 성장, 신규 고객 47개사 확보, 운영 비용 8% 감소."
    edits = [
        {
            "edit_type": "text",
            "target_kind": "paragraph",
            "target_id": target_id,
            "expected_text_hash": expected_hash,
            "new_text": new_text,
            "reason": "smoke test fill",
        }
    ]

    print(f"[4/5] fill_and_save -> {output_path}")
    apply_result = fill_and_save(
        template_path=str(template_path),
        edits=edits,
        output_path=str(output_path),
        dry_run=False,
    )
    print(json.dumps(apply_result, ensure_ascii=False, indent=2)[:1200])

    if apply_result.get("status") != "ok":
        print("FAIL: fill_and_save did not return ok")
        return 1
    if not output_path.is_file():
        print("FAIL: output file was not written")
        return 1

    print(f"[5/5] verify output contains new text")
    out_doc = Document(output_path)
    body_text = "\n".join(p.text for p in out_doc.paragraphs)
    print("--- output body ---")
    print(body_text)
    print("--- end body ---")
    if new_text not in body_text:
        print(f"FAIL: new_text not found in output document")
        return 1

    print("PASS: basic end-to-end smoke test succeeded")

    # ----------------------------------------------------------------------
    # TEST 2: style preservation
    # ----------------------------------------------------------------------
    print()
    print("=" * 70)
    print("TEST 2: verify original template styles survive an edit")
    print("=" * 70)

    styled_tmpl = tmp / "styled_template.docx"
    styled_out = tmp / "styled_filled.docx"
    sdoc = Document()
    sdoc.add_heading("회사 분기 보고서", level=0)

    p1 = sdoc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("핵심 메시지: 이번 분기 실적 요약")
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p2 = sdoc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("작성일: 2026-05-15")
    r2.italic = True
    r2.font.size = Pt(10)

    sdoc.save(styled_tmpl)

    # Snapshot original styles by content
    pre = Document(styled_tmpl)
    pre_snapshot = {}
    for p in pre.paragraphs:
        if not p.runs:
            continue
        r = p.runs[0]
        pre_snapshot[p.text] = {
            "alignment": p.alignment,
            "bold": r.bold,
            "italic": r.italic,
            "size": r.font.size,
            "color_rgb": str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None,
        }
    print("BEFORE edit — captured styles:")
    for k, v in pre_snapshot.items():
        print(f"  {k!r:60s} -> {v}")

    # Find the heading-styled paragraph and replace its text
    s_targets = list_template_targets(str(styled_tmpl), target_kinds=["paragraph"], max_targets=50)
    headline = None
    for t in s_targets.get("targets", []):
        if (t.get("current_text") or "").startswith("핵심 메시지"):
            headline = t
            break
    if headline is None:
        print("FAIL: could not find headline paragraph in styled template")
        return 1

    s_edits = [
        {
            "edit_type": "text",
            "target_kind": "paragraph",
            "target_id": headline["target_id"],
            "expected_text_hash": headline["text_hash"],
            "new_text": "핵심 메시지: 매출 +12%, 영업이익 +18%, 신규 고객 47개사",
            "reason": "fill headline with actual chatbot summary",
        }
    ]
    s_apply = fill_and_save(
        template_path=str(styled_tmpl),
        edits=s_edits,
        output_path=str(styled_out),
        dry_run=False,
    )
    if s_apply.get("status") != "ok":
        print(f"FAIL: styled fill_and_save -> {s_apply}")
        return 1

    post = Document(styled_out)
    post_snapshot = {}
    for p in post.paragraphs:
        if not p.runs:
            continue
        r = p.runs[0]
        post_snapshot[p.text] = {
            "alignment": p.alignment,
            "bold": r.bold,
            "italic": r.italic,
            "size": r.font.size,
            "color_rgb": str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None,
        }
    print()
    print("AFTER edit — actual styles:")
    for k, v in post_snapshot.items():
        print(f"  {k!r:60s} -> {v}")

    # Check the unchanged paragraph kept its style verbatim
    untouched_key = "작성일: 2026-05-15"
    if pre_snapshot.get(untouched_key) != post_snapshot.get(untouched_key):
        print(f"FAIL: untouched paragraph style changed")
        print(f"  before: {pre_snapshot.get(untouched_key)}")
        print(f"  after:  {post_snapshot.get(untouched_key)}")
        return 1

    # Check the edited paragraph kept ITS style (center align, bold, red, 16pt)
    pre_headline = pre_snapshot["핵심 메시지: 이번 분기 실적 요약"]
    edited_text = "핵심 메시지: 매출 +12%, 영업이익 +18%, 신규 고객 47개사"
    post_headline = post_snapshot.get(edited_text)
    if post_headline is None:
        print("FAIL: edited paragraph not found in output by new text")
        return 1
    style_checks = {
        "alignment": pre_headline["alignment"] == post_headline["alignment"],
        "bold": pre_headline["bold"] == post_headline["bold"],
        "font_size": pre_headline["size"] == post_headline["size"],
        "font_color": pre_headline["color_rgb"] == post_headline["color_rgb"],
    }
    print()
    print("Style-preservation checks on edited paragraph:")
    for name, ok in style_checks.items():
        print(f"  {name:15s}: {'OK' if ok else 'FAIL'}  (before={pre_headline}, after={post_headline})")
    if not all(style_checks.values()):
        return 1

    print()
    print("PASS: edited paragraph keeps original alignment/bold/size/color;")
    print("      untouched paragraph keeps its original style verbatim.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
