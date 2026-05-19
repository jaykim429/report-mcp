"""Second-round probes after patches.

  - describe_template on HWPX and DOCX
  - empty new_text as a deletion mechanism
  - sequential fill_and_save calls on the same source (idempotency)
"""

from __future__ import annotations

import shutil
import sys
import tempfile
from pathlib import Path

from docx import Document

from report_mcp.server import (
    describe_template,
    fill_and_save,
    list_template_targets,
)

P = Path(__file__).parent
PROJECT_ROOT = P.parent
HWPX = next((PROJECT_ROOT / "output" / "templates").glob("*.hwpx"))


def make_simple_docx(path: Path) -> None:
    d = Document()
    d.add_heading("Plain Doc", 1)
    d.add_paragraph("Intro paragraph.")
    d.add_paragraph("Second paragraph.")
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "A"
    t.rows[0].cells[1].text = "B"
    t.rows[1].cells[0].text = "1"
    t.rows[1].cells[1].text = "2"
    d.save(path)


def main() -> int:
    failures: list[str] = []
    tmp = Path(tempfile.mkdtemp(prefix="rmcp_r2_"))

    # describe_template — HWPX
    print("[probe] describe_template on HWPX")
    r = describe_template(str(HWPX))
    print(f"  status={r.get('status')}  format={r.get('source_doc_type')}")
    print(f"  total_paragraphs={r.get('total_paragraphs')}  pages={r.get('page_count')}")
    print(f"  has_tables={r.get('has_tables')}  has_images={r.get('has_images')}")
    print(f"  target_counts={r.get('target_counts')}")
    print(f"  top_paragraphs[0..2]: {r.get('top_paragraphs', [])[:3]}")
    for k in ("source_doc_type", "total_paragraphs", "target_counts", "top_paragraphs"):
        if r.get(k) in (None, [], {}):
            failures.append(f"describe(HWPX) missing or empty {k}")

    # describe_template — DOCX
    print("\n[probe] describe_template on DOCX")
    docx_p = tmp / "plain.docx"
    make_simple_docx(docx_p)
    r2 = describe_template(str(docx_p))
    print(f"  status={r2.get('status')}  format={r2.get('source_doc_type')}")
    print(f"  total_paragraphs={r2.get('total_paragraphs')}")
    print(f"  has_tables={r2.get('has_tables')}  has_images={r2.get('has_images')}")
    print(f"  top_paragraphs: {r2.get('top_paragraphs', [])}")
    if r2.get("status") != "ok":
        failures.append(f"describe(DOCX) status={r2.get('status')}")
    if r2.get("has_tables") is not True:
        failures.append("describe(DOCX) failed to detect table")

    # describe_template — missing file
    print("\n[probe] describe_template on missing path")
    r3 = describe_template("c:/no/such/file.docx")
    if r3.get("status") != "not_found":
        failures.append(f"describe(missing) wrong status: {r3.get('status')}")
    print(f"  status={r3.get('status')}  hint={r3.get('recovery_hint', '')[:60]}...")

    # empty new_text (treat as deletion of the text in a paragraph)
    print("\n[probe] empty new_text — paragraph text becomes blank?")
    src = tmp / "del_in.hwpx"
    shutil.copy2(HWPX, src)
    out = tmp / "del_out.hwpx"
    paras = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")
    r4 = fill_and_save(str(src), [{
        "edit_type": "text", "target_kind": "paragraph",
        "target_id": date_t["target_id"], "expected_text_hash": date_t["text_hash"],
        "new_text": "",
    }], str(out))
    print(f"  status={r4.get('status')}")
    if r4.get("status") == "ok":
        post = list_template_targets(str(out), target_kinds=["paragraph"], max_targets=400).get("targets", [])
        # The date paragraph should now be empty
        date_after = next((t for t in post if t["target_id"] == date_t["target_id"]), None)
        if date_after is None:
            print("  (target_id no longer present — text deletion may have removed paragraph)")
        else:
            print(f"  date paragraph text now: {(date_after.get('current_text') or '')!r}")
    else:
        print(f"  refused: {r4.get('recovery_hint')}")

    # sequential calls — first edit, then second edit on the output
    print("\n[probe] sequential fill_and_save calls")
    src2 = tmp / "seq1.hwpx"
    shutil.copy2(HWPX, src2)
    mid = tmp / "seq2.hwpx"
    final = tmp / "seq3.hwpx"

    p2 = list_template_targets(str(src2), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    d2 = next(t for t in p2 if (t.get("current_text") or "").strip() == "2026. 03. 23.")
    r5a = fill_and_save(str(src2), [{
        "edit_type": "text", "target_kind": "paragraph",
        "target_id": d2["target_id"], "expected_text_hash": d2["text_hash"],
        "new_text": "2026. 05. 17.",
    }], str(mid))
    print(f"  call 1: status={r5a.get('status')}")
    if r5a.get("status") != "ok":
        failures.append(f"sequential call 1 failed: {r5a}")

    p3 = list_template_targets(str(mid), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    d3 = next((t for t in p3 if (t.get("current_text") or "").strip() == "2026. 05. 17."), None)
    if d3 is None:
        failures.append("sequential: edited paragraph not findable on output")
    else:
        r5b = fill_and_save(str(mid), [{
            "edit_type": "text", "target_kind": "paragraph",
            "target_id": d3["target_id"], "expected_text_hash": d3["text_hash"],
            "new_text": "2026. 05. 18.",
        }], str(final))
        print(f"  call 2: status={r5b.get('status')}")
        if r5b.get("status") != "ok":
            failures.append(f"sequential call 2 failed: {r5b}")
        else:
            p4 = list_template_targets(str(final), target_kinds=["paragraph"], max_targets=400).get("targets", [])
            if not any((t.get("current_text") or "").strip() == "2026. 05. 18." for t in p4):
                failures.append("sequential: final text not landed")

    print()
    if failures:
        print(f"FAIL ({len(failures)} issues):")
        for f in failures:
            print(f"  - {f}")
        return 1
    print("PASS — round-two probes all green")
    return 0


if __name__ == "__main__":
    sys.exit(main())
