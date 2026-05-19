"""Probe likely-broken edge cases in the report-mcp server.

Each probe prints PASS / FAIL / WARN and a short note. Collect all findings,
then patch server.py based on what's actually broken vs already handled.
"""

from __future__ import annotations

import shutil
import sys
import tempfile
import time
from pathlib import Path

from docx import Document
from docx.shared import Pt

from report_mcp.server import (
    fill_and_save,
    inspect_template,
    list_template_targets,
)

P = Path(__file__).parent
PROJECT_ROOT = P.parent
HWPX = next((PROJECT_ROOT / "output" / "templates").glob("*.hwpx"))

findings: list[tuple[str, str, str]] = []  # (status, probe_id, note)


def record(status: str, probe_id: str, note: str = "") -> None:
    findings.append((status, probe_id, note))
    print(f"  {status:4s} {probe_id}: {note}")


def make_docx(path: Path) -> None:
    d = Document()
    d.add_heading("월간 운영 보고서", level=1)
    d.add_paragraph("작성자: [AUTHOR]")
    d.add_paragraph("작성일: [DATE]")
    d.add_heading("1. 주요 성과", level=2)
    d.add_paragraph("[ACHIEVEMENT_1]")
    d.add_paragraph("[ACHIEVEMENT_2]")
    d.add_heading("2. 향후 계획", level=2)
    d.add_paragraph("[PLAN_1]")
    table = d.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "항목"
    table.rows[0].cells[1].text = "값"
    table.rows[1].cells[0].text = "처리량"
    table.rows[1].cells[1].text = "[METRIC_1]"
    table.rows[2].cells[0].text = "가용성"
    table.rows[2].cells[1].text = "[METRIC_2]"
    d.save(path)


def probe_docx_round_trip(workdir: Path) -> None:
    print("\n[probe] DOCX input round-trip")
    src = workdir / "monthly.docx"
    out = workdir / "monthly_filled.docx"
    make_docx(src)

    tgts = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=400)
    if tgts.get("status") != "ok":
        record("FAIL", "docx.list", f"list_template_targets returned {tgts.get('status')}")
        return
    paras = tgts.get("targets", [])
    record("PASS", "docx.list", f"{len(paras)} paragraph targets")

    author = next((t for t in paras if (t.get("current_text") or "").endswith("[AUTHOR]")), None)
    plan1 = next((t for t in paras if (t.get("current_text") or "") == "[PLAN_1]"), None)
    if not (author and plan1):
        record("WARN", "docx.locate", "could not find expected [AUTHOR]/[PLAN_1] anchors")
        return

    edits = [
        {"edit_type": "text", "target_kind": "paragraph",
         "target_id": author["target_id"], "expected_text_hash": author["text_hash"],
         "new_text": "작성자: 김정훈"},
        {"edit_type": "text", "target_kind": "paragraph",
         "target_id": plan1["target_id"], "expected_text_hash": plan1["text_hash"],
         "new_text": "AI 챗봇 도입 PoC 본격 실행"},
    ]
    r = fill_and_save(str(src), edits, str(out))
    if r.get("status") != "ok":
        record("FAIL", "docx.fill", f"{r}")
        return
    if not out.is_file():
        record("FAIL", "docx.fill", "output file missing")
        return

    # Read back via python-docx and assert
    post = Document(out)
    bodies = [p.text for p in post.paragraphs]
    if "작성자: 김정훈" not in bodies:
        record("FAIL", "docx.verify", "[AUTHOR] not replaced in output")
    elif "AI 챗봇 도입 PoC 본격 실행" not in bodies:
        record("FAIL", "docx.verify", "[PLAN_1] not replaced in output")
    else:
        record("PASS", "docx.verify", "both edits visible in output DOCX")


def probe_empty_edits(workdir: Path) -> None:
    print("\n[probe] empty edits list")
    src = workdir / "empty_in.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "empty_out.hwpx"
    r = fill_and_save(str(src), [], str(out))
    if r.get("status") == "ok":
        # Does the output exist? Should be a copy of input.
        if out.is_file():
            record("PASS", "empty.ok", "empty edit list returns ok and produces file")
        else:
            record("FAIL", "empty.nofile", "ok but no output file written")
    else:
        record("WARN", "empty.status", f"empty edits returned status={r.get('status')}, would be cleaner to short-circuit to ok with note")


def probe_duplicate_target(workdir: Path) -> None:
    print("\n[probe] duplicate target_id in edits")
    src = workdir / "dup_in.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "dup_out.hwpx"

    paras = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")
    e1 = {"edit_type": "text", "target_kind": "paragraph",
          "target_id": date_t["target_id"], "expected_text_hash": date_t["text_hash"],
          "new_text": "2026. 05. 16."}
    e2 = {"edit_type": "text", "target_kind": "paragraph",
          "target_id": date_t["target_id"], "expected_text_hash": date_t["text_hash"],
          "new_text": "2026. 05. 18."}
    r = fill_and_save(str(src), [e1, e2], str(out))
    if r.get("status") != "ok":
        record("WARN", "dup.fails", f"two edits on same target rejected: {r.get('status')} -> {r.get('recovery_hint') or r.get('validation', {})}")
    else:
        # Which one wins?
        post = list_template_targets(str(out), target_kinds=["paragraph"], max_targets=400).get("targets", [])
        match = next((t for t in post if (t.get("current_text") or "").strip().startswith("2026. 05.")), None)
        if match:
            record("PASS", "dup.last_wins", f"final text = {match.get('current_text')!r}")
        else:
            record("FAIL", "dup.lost", "neither edit visible")


def probe_unknown_target(workdir: Path) -> None:
    print("\n[probe] unknown target_id")
    src = workdir / "unk_in.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "unk_out.hwpx"
    r = fill_and_save(str(src), [{
        "edit_type": "text", "target_kind": "paragraph",
        "target_id": "p_nonexistent_99999", "expected_text_hash": "0" * 40,
        "new_text": "nope",
    }], str(out))
    if r.get("status") == "ok":
        record("FAIL", "unk.ok", "expected validation failure for fake target_id")
    elif "recovery_hint" in r:
        record("PASS", "unk.fails_with_hint", f"status={r.get('status')}, hint includes 'list_template_targets'")
    else:
        record("WARN", "unk.no_hint", f"status={r.get('status')} but no recovery_hint")


def probe_output_overwrites_source(workdir: Path) -> None:
    print("\n[probe] output_path equals template_path (source overwrite)")
    src = workdir / "self.hwpx"
    shutil.copy2(HWPX, src)
    paras = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")
    edits = [{"edit_type": "text", "target_kind": "paragraph",
              "target_id": date_t["target_id"], "expected_text_hash": date_t["text_hash"],
              "new_text": "2026. 05. 18."}]
    r = fill_and_save(str(src), edits, str(src))
    if r.get("status") == "ok":
        # Verify edit landed
        post = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=400).get("targets", [])
        if any((t.get("current_text") or "").strip() == "2026. 05. 18." for t in post):
            record("PASS", "self.replace", "in-place overwrite works")
        else:
            record("FAIL", "self.lost", "ok status but edit not in result file")
    else:
        record("WARN", "self.refused", f"in-place overwrite rejected: {r.get('status')}")


def probe_dry_run(workdir: Path) -> None:
    print("\n[probe] dry_run")
    src = workdir / "dry_in.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "dry_out.hwpx"
    paras = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")
    r = fill_and_save(str(src), [{
        "edit_type": "text", "target_kind": "paragraph",
        "target_id": date_t["target_id"], "expected_text_hash": date_t["text_hash"],
        "new_text": "2026. 05. 18.",
    }], str(out), dry_run=True)
    if r.get("status") != "dry_run_ok":
        record("FAIL", "dry.status", f"expected dry_run_ok, got {r.get('status')}")
    elif out.exists():
        record("FAIL", "dry.wrote", "dry_run should NOT write output file")
    else:
        record("PASS", "dry.clean", f"dry_run_ok with batches_planned={r.get('batches_planned')}, no file written")


def probe_hash_stability(workdir: Path) -> None:
    print("\n[probe] hash stability across repeated list calls")
    a = list_template_targets(str(HWPX), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    time.sleep(0.05)
    b = list_template_targets(str(HWPX), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    drift = sum(1 for x, y in zip(a, b) if x.get("text_hash") != y.get("text_hash"))
    if drift == 0:
        record("PASS", "hash.stable", f"{len(a)} targets, 0 hash drift")
    else:
        record("FAIL", "hash.drift", f"{drift} of {len(a)} targets changed hash between calls")


def probe_overlong_single(workdir: Path) -> None:
    print("\n[probe] single very-long replacement (way past max_recommended)")
    src = workdir / "long_in.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "long_out.hwpx"
    paras = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")
    r = fill_and_save(str(src), [{
        "edit_type": "text", "target_kind": "paragraph",
        "target_id": date_t["target_id"], "expected_text_hash": date_t["text_hash"],
        "new_text": "엄청 긴 텍스트 " * 100,  # ~700 chars
    }], str(out))
    if r.get("status") != "ok":
        record("FAIL", "long.status", f"unexpected non-ok: {r.get('status')}")
    elif not r.get("length_warnings"):
        record("FAIL", "long.nowarn", "expected length warning but got none")
    else:
        w = r["length_warnings"][0]
        record("PASS", "long.warn",
               f"warned: {w['original_display_width']}->{w['new_display_width']} cells "
               f"(cap {w['max_recommended_width']}, +{w['overflow_cells']})")


def probe_run_edit(workdir: Path) -> None:
    print("\n[probe] TextEdit on run (inline) target")
    src = workdir / "run_in.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "run_out.hwpx"
    runs = list_template_targets(str(src), target_kinds=["run"], max_targets=400).get("targets", [])
    if not runs:
        record("WARN", "run.none", "no run targets returned for HWPX — chatbot won't use this")
        return
    r0 = next((r for r in runs if (r.get("current_text") or "").strip()), None)
    if r0 is None:
        record("WARN", "run.empty", "all run targets had empty text")
        return
    edits = [{"edit_type": "text", "target_kind": "run",
              "target_id": r0["target_id"], "expected_text_hash": r0["text_hash"],
              "new_text": "교체된 인라인 텍스트"}]
    r = fill_and_save(str(src), edits, str(out))
    if r.get("status") == "ok":
        record("PASS", "run.ok", f"run-level edit applied")
    else:
        record("WARN", "run.fails", f"status={r.get('status')}")


def probe_missing_output_dir(workdir: Path) -> None:
    print("\n[probe] output_path under non-existent directory")
    src = workdir / "mkdir_in.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "deep" / "nested" / "new" / "out.hwpx"
    paras = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")
    r = fill_and_save(str(src), [{
        "edit_type": "text", "target_kind": "paragraph",
        "target_id": date_t["target_id"], "expected_text_hash": date_t["text_hash"],
        "new_text": "2026. 05. 18.",
    }], str(out))
    if r.get("status") == "ok" and out.is_file():
        record("PASS", "mkdir.auto", "deeply nested missing dirs were created")
    else:
        record("FAIL", "mkdir.fail", f"status={r.get('status')}, file={out.is_file()}")


def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="rmcp_probe_"))
    print(f"working dir: {tmp}\n")

    probe_docx_round_trip(tmp)
    probe_empty_edits(tmp)
    probe_duplicate_target(tmp)
    probe_unknown_target(tmp)
    probe_output_overwrites_source(tmp)
    probe_dry_run(tmp)
    probe_hash_stability(tmp)
    probe_overlong_single(tmp)
    probe_run_edit(tmp)
    probe_missing_output_dir(tmp)

    print("\n" + "=" * 70)
    fails = [f for f in findings if f[0] == "FAIL"]
    warns = [f for f in findings if f[0] == "WARN"]
    passes = [f for f in findings if f[0] == "PASS"]
    print(f"PASS={len(passes)}  WARN={len(warns)}  FAIL={len(fails)}")
    for s, pid, note in findings:
        print(f"  {s:4s} {pid}: {note}")
    return 1 if fails else 0


if __name__ == "__main__":
    sys.exit(main())
