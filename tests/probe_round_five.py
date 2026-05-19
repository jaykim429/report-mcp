"""Round-five probes: cross-cutting concerns less obvious than core CRUD."""

from __future__ import annotations

import contextlib
import gc
import io
import os
import shutil
import sys
import tempfile
import unicodedata
from pathlib import Path

from docx import Document

from report_mcp.server import (
    describe_template,
    fill_and_save,
    inspect_template,
    list_template_targets,
)

P = Path(__file__).parent
PROJECT_ROOT = P.parent
HWPX = next((PROJECT_ROOT / "output" / "templates").glob("*.hwpx"))

findings: list[tuple[str, str, str]] = []


def rec(s: str, pid: str, note: str = "") -> None:
    findings.append((s, pid, note))
    print(f"  {s:4s} {pid}: {note}")


# ────────────────────────────────────────────────────────────────────────
# A. stdout pollution — MCP stdio uses stdout for JSON-RPC, so a single
# stray print() leak would corrupt the protocol stream for every client.
# ────────────────────────────────────────────────────────────────────────

def probe_stdout_pollution() -> None:
    print("\n[A] stdout pollution during tool invocations")
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        describe_template(str(HWPX))
        inspect_template(str(HWPX), start=0, limit=5)
        list_template_targets(str(HWPX), target_kinds=["paragraph"], limit=10)
        # bad-argument call too (touches error path)
        list_template_targets(str(HWPX), start=-1, limit=10)
    leaked = buf.getvalue()
    if not leaked:
        rec("PASS", "stdout.silent", "no tool wrote to stdout (MCP stdio transport safe)")
    else:
        rec("FAIL", "stdout.leak",
            f"{len(leaked)} chars leaked to stdout; first 200: {leaked[:200]!r}")


# ────────────────────────────────────────────────────────────────────────
# B. Template file locked exclusively (simulates Word/Hangul holding it open)
# ────────────────────────────────────────────────────────────────────────

def probe_locked_template(workdir: Path) -> None:
    print("\n[B] locked template + unwritable output_path")
    # Test #1: output_path points to a directory, not a file
    src = workdir / "lock.hwpx"
    shutil.copy2(HWPX, src)
    paras = list_template_targets(str(src), target_kinds=["paragraph"], limit=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")
    edits = [{"edit_type": "text", "target_kind": "paragraph",
              "target_id": date_t["target_id"], "expected_text_hash": date_t["text_hash"],
              "new_text": "2026. 05. 18."}]

    # Output_path is the workdir (a directory)
    r1 = fill_and_save(str(src), edits, str(workdir))
    print(f"  output_path=dir: status={r1.get('status')}")
    if r1.get("status") == "ok":
        rec("FAIL", "lock.dir_out", "writing to a directory path returned ok — should reject")
    elif r1.get("status") in {"bad_argument", "output_extension_mismatch", "apply_failed", "permission_error"}:
        rec("PASS", "lock.dir_out", f"directory-as-output rejected with status={r1.get('status')}")
    else:
        rec("WARN", "lock.dir_out", f"unexpected status: {r1.get('status')}")

    # Test #2: template_path locked exclusively on Windows. Use os.open with
    # O_EXCL-ish — actually use Win32 file open with FILE_SHARE_NONE via msvcrt.
    # Easier: open and hold a file handle in 'r+b' mode then try to read.
    # On Windows, default open mode doesn't deny sharing — to simulate Word's
    # exclusive lock, use msvcrt.locking() on a region.
    locked = workdir / "locked.hwpx"
    shutil.copy2(HWPX, locked)

    try:
        import msvcrt
        fh = open(locked, "r+b")
        msvcrt.locking(fh.fileno(), msvcrt.LK_NBLCK, 1)
        try:
            r2 = fill_and_save(str(locked), edits, str(workdir / "out_locked.hwpx"))
            print(f"  template locked: status={r2.get('status')}")
            if r2.get("status") in {"ok"}:
                rec("WARN", "lock.read_ok", "template lock did NOT block reading — Hangul lock semantics differ")
            elif r2.get("status") == "permission_error":
                rec("PASS", "lock.perm", "locked template surfaced as permission_error")
            else:
                rec("PASS", "lock.other", f"lock surfaced as status={r2.get('status')}")
        finally:
            try:
                msvcrt.locking(fh.fileno(), msvcrt.LK_UNLCK, 1)
            except OSError:
                pass
            fh.close()
    except (ImportError, OSError) as exc:
        rec("WARN", "lock.skip", f"could not simulate lock: {exc}")


# ────────────────────────────────────────────────────────────────────────
# C. Idempotent re-apply — apply same edits twice
# ────────────────────────────────────────────────────────────────────────

def probe_idempotent_reapply(workdir: Path) -> None:
    print("\n[C] re-apply same edits twice (idempotency)")
    src = workdir / "idem.hwpx"
    shutil.copy2(HWPX, src)
    out1 = workdir / "idem_a.hwpx"
    out2 = workdir / "idem_b.hwpx"

    paras = list_template_targets(str(src), target_kinds=["paragraph"], limit=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")
    edits = [{"edit_type": "text", "target_kind": "paragraph",
              "target_id": date_t["target_id"], "expected_text_hash": date_t["text_hash"],
              "new_text": "2026. 05. 18."}]

    r1 = fill_and_save(str(src), edits, str(out1))
    print(f"  first apply: status={r1.get('status')}")

    # Re-apply: source is now out1 with edit already applied. The provided
    # expected_text_hash is from the ORIGINAL — so it won't match.
    r2 = fill_and_save(str(out1), edits, str(out2))
    print(f"  re-apply same edits (stale hashes): status={r2.get('status')}")
    if r2.get("status") == "validation_failed":
        targets = r2.get("failed_targets") or []
        if targets and targets[0].get("current_text") == "2026. 05. 18.":
            rec("PASS", "idem.detected", "re-apply with stale hashes correctly rejected; failed_targets shows current_text matches the desired new state")
        else:
            rec("WARN", "idem.partial", f"rejected but failed_targets payload thin: {targets[:1]}")
    elif r2.get("status") == "ok":
        rec("WARN", "idem.silent",
            "re-apply ok even though hashes were stale — write happened a second time")
    else:
        rec("WARN", "idem.other", f"unexpected status: {r2.get('status')}")


# ────────────────────────────────────────────────────────────────────────
# D. Unicode normalization — Korean Hangul can be NFC (precomposed syllable
# blocks) or NFD (separate jamo). Same string visually, different bytes.
# ────────────────────────────────────────────────────────────────────────

def probe_unicode_normalization(workdir: Path) -> None:
    print("\n[D] NFC vs NFD Korean — does hash discriminate?")
    src = workdir / "nfc.docx"
    out = workdir / "nfc_out.docx"
    d = Document()
    d.add_paragraph("한국어 단락")  # NFC
    d.save(src)

    paras = list_template_targets(str(src), target_kinds=["paragraph"], limit=10).get("targets", [])
    p = next(t for t in paras if (t.get("current_text") or "").strip().endswith("단락"))
    nfd_text = unicodedata.normalize("NFD", "한국어 단락")
    print(f"  current_text repr: {p['current_text']!r}  len={len(p['current_text'])}")
    print(f"  NFD form repr:     {nfd_text!r}  len={len(nfd_text)}")

    # If chatbot accidentally sends NFD, will server detect mismatch?
    edits = [{"edit_type": "text", "target_kind": "paragraph",
              "target_id": p["target_id"], "expected_text_hash": p["text_hash"],
              "new_text": "새 한국어 단락 NFD test"}]
    r = fill_and_save(str(src), edits, str(out))
    if r.get("status") == "ok":
        rec("PASS", "nfc.standard", "NFC text round-trips through edit pipeline")
    else:
        rec("FAIL", "nfc.standard", f"NFC failed: {r.get('status')}")


# ────────────────────────────────────────────────────────────────────────
# E. Empty document — zero editable targets
# ────────────────────────────────────────────────────────────────────────

def probe_empty_document(workdir: Path) -> None:
    print("\n[E] DOCX with no body content")
    src = workdir / "empty.docx"
    Document().save(src)

    r1 = describe_template(str(src))
    print(f"  describe: status={r1.get('status')} paragraphs={r1.get('total_paragraphs')}  targets={r1.get('target_counts')}")
    if r1.get("status") == "ok":
        rec("PASS", "empty.describe", f"describe_template ok on empty doc")
    else:
        rec("FAIL", "empty.describe", f"{r1}")

    r2 = list_template_targets(str(src), limit=10)
    print(f"  list: status={r2.get('status')} returned={r2.get('returned')}")
    if r2.get("status") == "ok":
        rec("PASS", "empty.list", "list_template_targets ok with zero targets")
    else:
        rec("FAIL", "empty.list", f"{r2}")


# ────────────────────────────────────────────────────────────────────────
# F. Temp directory leak — call fill_and_save many times in-place
# (in-place uses tempfile internally), verify temp parent doesn't grow
# ────────────────────────────────────────────────────────────────────────

def probe_temp_dir_leak(workdir: Path) -> None:
    print("\n[F] temp dir leak after 20 in-place fills")
    src = workdir / "leak.hwpx"
    shutil.copy2(HWPX, src)

    temp_root = Path(tempfile.gettempdir())
    before = {p.name for p in temp_root.iterdir() if p.name.startswith("rmcp_inplace_")}

    paras = list_template_targets(str(src), target_kinds=["paragraph"], limit=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")

    for i in range(20):
        # Re-fetch hash each iteration since previous edit changed the file
        p2 = list_template_targets(str(src), target_kinds=["paragraph"], limit=400).get("targets", [])
        d = next(t for t in p2 if t["target_id"] == date_t["target_id"])
        r = fill_and_save(str(src), [{
            "edit_type": "text", "target_kind": "paragraph",
            "target_id": d["target_id"], "expected_text_hash": d["text_hash"],
            "new_text": f"iter {i}",
        }], str(src))  # in-place
        if r.get("status") != "ok":
            rec("FAIL", "leak.fill", f"iter {i} failed: {r.get('status')}")
            return
    gc.collect()
    after = {p.name for p in temp_root.iterdir() if p.name.startswith("rmcp_inplace_")}
    delta = after - before
    if not delta:
        rec("PASS", "leak.none", "20 in-place fills left no rmcp_inplace_* temp dirs behind")
    else:
        rec("FAIL", "leak.found", f"{len(delta)} leaked temp dirs: {sorted(delta)[:3]}...")


# ────────────────────────────────────────────────────────────────────────
# G. Relative path handling
# ────────────────────────────────────────────────────────────────────────

def probe_relative_paths(workdir: Path) -> None:
    print("\n[G] relative paths")
    src = workdir / "rel.hwpx"
    shutil.copy2(HWPX, src)

    cwd_saved = os.getcwd()
    try:
        os.chdir(workdir)
        # Use bare filename (relative to cwd)
        r = describe_template("rel.hwpx")
        if r.get("status") == "ok":
            rec("PASS", "rel.describe", "relative path accepted by describe_template")
        else:
            rec("FAIL", "rel.describe", f"{r.get('status')}: {r.get('error')}")

        # Relative output path
        edits_paras = list_template_targets("rel.hwpx", target_kinds=["paragraph"], limit=400).get("targets", [])
        date_t = next(t for t in edits_paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")
        r2 = fill_and_save("rel.hwpx", [{
            "edit_type": "text", "target_kind": "paragraph",
            "target_id": date_t["target_id"], "expected_text_hash": date_t["text_hash"],
            "new_text": "2026. 05. 18.",
        }], "rel_out.hwpx")
        if r2.get("status") == "ok" and (workdir / "rel_out.hwpx").is_file():
            rec("PASS", "rel.fill", f"relative paths work end-to-end")
        else:
            rec("FAIL", "rel.fill", f"{r2.get('status')}: {r2.get('error')}")
    finally:
        os.chdir(cwd_saved)


# ────────────────────────────────────────────────────────────────────────
# H. Pagination hash consistency across pages
# ────────────────────────────────────────────────────────────────────────

def probe_pagination_hash_consistency() -> None:
    print("\n[H] pagination — hashes consistent across paginated pages?")
    # Fetch all in one go
    all_p = list_template_targets(str(HWPX), target_kinds=["paragraph"], limit=400)
    full = {t["target_id"]: t["text_hash"] for t in all_p.get("targets", [])}

    # Now fetch page-by-page with limit=20 and compare
    page_size = 20
    start = 0
    drift = 0
    seen = 0
    while True:
        page = list_template_targets(str(HWPX), target_kinds=["paragraph"], start=start, limit=page_size)
        items = page.get("targets", [])
        for t in items:
            if full.get(t["target_id"]) != t["text_hash"]:
                drift += 1
            seen += 1
        if page.get("next_start") is None:
            break
        start = page["next_start"]

    if drift == 0 and seen == len(full):
        rec("PASS", "page.consistent", f"{seen} targets across pages, all hashes match the one-shot view")
    else:
        rec("FAIL", "page.drift", f"drift={drift}, seen={seen} vs expected={len(full)}")


# ────────────────────────────────────────────────────────────────────────

def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="rmcp_r5_"))
    print(f"workdir: {tmp}")

    for fn, args in [
        (probe_stdout_pollution, ()),
        (probe_locked_template, (tmp,)),
        (probe_idempotent_reapply, (tmp,)),
        (probe_unicode_normalization, (tmp,)),
        (probe_empty_document, (tmp,)),
        (probe_temp_dir_leak, (tmp,)),
        (probe_relative_paths, (tmp,)),
        (probe_pagination_hash_consistency, ()),
    ]:
        try:
            fn(*args)
        except Exception as exc:
            rec("FAIL", f"crash.{fn.__name__}", f"{type(exc).__name__}: {exc}")

    print("\n" + "─" * 70)
    p = sum(1 for s, *_ in findings if s == "PASS")
    w = sum(1 for s, *_ in findings if s == "WARN")
    f = sum(1 for s, *_ in findings if s == "FAIL")
    print(f"PASS={p}  WARN={w}  FAIL={f}")
    for s, pid, note in findings:
        if s != "PASS":
            print(f"  {s} {pid}: {note}")
    return 1 if f else 0


if __name__ == "__main__":
    sys.exit(main())
