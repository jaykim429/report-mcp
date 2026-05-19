"""Round-three probes: StructuralEdit, StyleEdit, MCP boot integrity."""

from __future__ import annotations

import asyncio
import shutil
import sys
import tempfile
from pathlib import Path

from docx import Document

from report_mcp.server import fill_and_save, list_template_targets, mcp

P = Path(__file__).parent
PROJECT_ROOT = P.parent
HWPX = next((PROJECT_ROOT / "output" / "templates").glob("*.hwpx"))

results: list[tuple[str, str, str]] = []


def record(status: str, pid: str, note: str = "") -> None:
    results.append((status, pid, note))
    print(f"  {status:4s} {pid}: {note}")


# ---------------------------------------------------------------------------

def probe_structural_insert_paragraph(workdir: Path) -> None:
    """StructuralEdit: insert a paragraph after a known anchor."""
    print("\n[probe] StructuralEdit.insert_paragraph")
    src = workdir / "ins_p.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "ins_p_out.hwpx"

    paras = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")

    edits = [{
        "edit_type": "structural",
        "operation": "insert_paragraph",
        "target_id": date_t["target_id"],
        "position": "after",
        "text": "[새로 삽입된 단락]",
    }]
    r = fill_and_save(str(src), edits, str(out))
    if r.get("status") != "ok":
        record("FAIL", "ins_p.status", f"{r.get('status')} - {r.get('error') or r.get('recovery_hint')}")
        return

    post = list_template_targets(str(out), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    if any((t.get("current_text") or "").strip() == "[새로 삽입된 단락]" for t in post):
        record("PASS", "ins_p.found", f"new paragraph visible; total grew from {len(paras)} to {len(post)}")
    else:
        record("FAIL", "ins_p.lost", "new paragraph text not found after insert")


def probe_structural_remove_paragraph(workdir: Path) -> None:
    """StructuralEdit: remove a paragraph."""
    print("\n[probe] StructuralEdit.remove_paragraph")
    src = workdir / "rm_p.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "rm_p_out.hwpx"

    paras = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    date_t = next(t for t in paras if (t.get("current_text") or "").strip() == "2026. 03. 23.")

    edits = [{
        "edit_type": "structural",
        "operation": "remove_paragraph",
        "target_id": date_t["target_id"],
    }]
    r = fill_and_save(str(src), edits, str(out))
    if r.get("status") != "ok":
        record("WARN", "rm_p.status", f"{r.get('status')} - {r.get('error') or r.get('recovery_hint', '')[:80]}")
        return

    post = list_template_targets(str(out), target_kinds=["paragraph"], max_targets=400).get("targets", [])
    if not any((t.get("current_text") or "").strip() == "2026. 03. 23." for t in post):
        record("PASS", "rm_p.gone", f"target removed; total shrank from {len(paras)} to {len(post)}")
    else:
        record("FAIL", "rm_p.persist", "target paragraph still present after remove")


def probe_structural_set_cell_text(workdir: Path) -> None:
    """StructuralEdit.set_cell_text: must NOT be auto-filtered (it's the
    canonical way to set cell content as a whole, distinct from TextEdit-on-cell)."""
    print("\n[probe] StructuralEdit.set_cell_text (must NOT be filtered)")
    src = workdir / "scell.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "scell_out.hwpx"

    cells = list_template_targets(str(src), target_kinds=["cell"], max_targets=400).get("targets", [])
    target_cell = next((c for c in cells if (c.get("current_text") or "").strip() == "PoC 구성용 아키텍쳐"), None)
    if target_cell is None:
        record("WARN", "scell.skip", "no cell with that text — different fixture")
        return

    edits = [{
        "edit_type": "structural",
        "operation": "set_cell_text",
        "target_id": target_cell["target_id"],
        "text": "셀 단위 교체 텍스트",
    }]
    r = fill_and_save(str(src), edits, str(out))
    if r.get("status") != "ok":
        record("WARN", "scell.status", f"{r.get('status')} - hint: {(r.get('recovery_hint') or '')[:80]}")
        return

    if r.get("skipped_redundant_edits"):
        record("FAIL", "scell.filtered", "StructuralEdit on cell was wrongly auto-filtered")
        return

    post = list_template_targets(str(out), target_kinds=["cell"], max_targets=400).get("targets", [])
    new_text = next((c.get("current_text", "") for c in post if c["target_id"] == target_cell["target_id"]), "")
    if "셀 단위 교체 텍스트" in new_text:
        record("PASS", "scell.set", f"cell text replaced; current='{new_text[:30]}'")
    else:
        record("FAIL", "scell.lost", f"cell unchanged after set_cell_text: '{new_text[:30]}'")


def probe_structural_insert_table_row(workdir: Path) -> None:
    """StructuralEdit.insert_table_row on a DOCX with a known table."""
    print("\n[probe] StructuralEdit.insert_table_row (DOCX)")
    src = workdir / "ins_row.docx"
    out = workdir / "ins_row_out.docx"
    d = Document()
    d.add_heading("Pricing", 1)
    t = d.add_table(rows=3, cols=3)
    for i, label in enumerate(["항목", "수량", "금액"]):
        t.rows[0].cells[i].text = label
    t.rows[1].cells[0].text = "API 사용료"
    t.rows[1].cells[1].text = "1식"
    t.rows[1].cells[2].text = "10,000,000"
    t.rows[2].cells[0].text = "합계"
    t.rows[2].cells[2].text = "10,000,000"
    d.save(src)

    tables = list_template_targets(str(src), target_kinds=["table"], max_targets=10).get("targets", [])
    if not tables:
        record("WARN", "ins_row.notable", "no table targets — unexpected")
        return
    tbl = tables[0]
    edits = [{
        "edit_type": "structural",
        "operation": "insert_table_row",
        "target_id": tbl["target_id"],
        "row_index": 2,  # insert before the 합계 row
        "values": ["인건비", "8주", "20,000,000"],
    }]
    r = fill_and_save(str(src), edits, str(out))
    if r.get("status") != "ok":
        record("WARN", "ins_row.status", f"{r.get('status')} - {(r.get('recovery_hint') or '')[:80]}")
        return

    post = Document(out)
    row_count = len(post.tables[0].rows)
    row_texts = [[c.text for c in row.cells] for row in post.tables[0].rows]
    if row_count == 4 and any(c == "인건비" for r in row_texts for c in r):
        record("PASS", "ins_row.added", f"row count grew 3->4, new content visible")
    else:
        record("FAIL", "ins_row.bad", f"row count={row_count}, rows={row_texts}")


def probe_style_edit_bold(workdir: Path) -> None:
    """StyleEdit: paragraph-level fields work on paragraph target;
    run-level fields work on run target; mixing them returns a clean
    style_field_target_mismatch."""
    print("\n[probe] StyleEdit on paragraph + run + mismatched targeting")
    src = workdir / "style.docx"
    d = Document()
    d.add_paragraph("강조하고 싶은 한 줄 문장입니다.")
    d.save(src)

    paras = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=20).get("targets", [])
    runs  = list_template_targets(str(src), target_kinds=["run"], max_targets=20).get("targets", [])
    p0 = paras[0]
    r0 = runs[0] if runs else None

    # 1) Paragraph-level fields ON paragraph -> should succeed
    out1 = workdir / "style_para.docx"
    r1 = fill_and_save(str(src), [{
        "edit_type": "style", "target_kind": "paragraph",
        "target_id": p0["target_id"],
        "paragraph_align": "center",
    }], str(out1))
    if r1.get("status") == "ok":
        align_after = Document(out1).paragraphs[0].alignment
        record("PASS", "style.para_ok", f"paragraph_align landed -> {align_after}")
    else:
        record("FAIL", "style.para_fail", f"{r1.get('status')} {r1.get('error')}")

    # 2) Run-level fields ON paragraph -> should return clean mismatch hint
    out2 = workdir / "style_bad.docx"
    r2 = fill_and_save(str(src), [{
        "edit_type": "style", "target_kind": "paragraph",
        "target_id": p0["target_id"],
        "bold": True, "color": "#C00000", "font_size_pt": 18.0,
    }], str(out2))
    if r2.get("status") == "style_field_target_mismatch":
        record("PASS", "style.mismatch_hint", f"clean status + hint='{r2['recovery_hint'][:60]}...'")
    else:
        record("FAIL", "style.mismatch_hint", f"expected style_field_target_mismatch, got {r2.get('status')}")

    # 3) Run-level fields ON run -> should succeed
    if r0 is None:
        record("WARN", "style.run_skip", "no run target available")
        return
    out3 = workdir / "style_run.docx"
    r3 = fill_and_save(str(src), [{
        "edit_type": "style", "target_kind": "run",
        "target_id": r0["target_id"],
        "bold": True, "color": "#C00000", "font_size_pt": 18.0,
    }], str(out3))
    if r3.get("status") == "ok":
        runs_after = Document(out3).paragraphs[0].runs
        record("PASS", "style.run_ok",
               f"bold={runs_after[0].bold}, size={runs_after[0].font.size}, "
               f"color={runs_after[0].font.color.rgb if runs_after[0].font.color and runs_after[0].font.color.rgb else None}")
    else:
        record("FAIL", "style.run_fail", f"{r3.get('status')} {r3.get('error')}")


def probe_mcp_boot_integrity() -> None:
    """Verify mcp instance has all tools registered with valid JSON schemas
    that the MCP protocol clients can consume."""
    print("\n[probe] MCP boot integrity (tool registration + schemas)")
    # FastMCP exposes async list_tools(); the run() call would serve these.
    try:
        tools = asyncio.run(mcp.list_tools())
    except Exception as exc:
        record("FAIL", "boot.list_tools", f"{type(exc).__name__}: {exc}")
        return

    names = [t.name for t in tools]
    expected = {"inspect_template", "list_template_targets", "fill_and_save", "describe_template"}
    missing = expected - set(names)
    if missing:
        record("FAIL", "boot.missing", f"missing tools: {missing}")
        return

    record("PASS", "boot.tools", f"all 4 tools registered: {names}")

    # Check each tool has description + inputSchema
    bad = []
    for t in tools:
        if not getattr(t, "description", None):
            bad.append(f"{t.name}: no description")
        schema = getattr(t, "inputSchema", None)
        if not isinstance(schema, dict) or "properties" not in schema:
            bad.append(f"{t.name}: malformed inputSchema")
    if bad:
        record("FAIL", "boot.schemas", "; ".join(bad))
    else:
        record("PASS", "boot.schemas", "every tool carries description + JSON schema")

    # Check server instructions visible
    instr = getattr(mcp, "instructions", None) or getattr(mcp._mcp_server, "instructions", None)
    if instr and "USE THIS SERVER WHEN" in instr:
        record("PASS", "boot.instructions", f"server instructions length={len(instr)}")
    else:
        record("FAIL", "boot.instructions", "server instructions not set or missing trigger guidance")


# ---------------------------------------------------------------------------

def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="rmcp_r3_"))
    print(f"workdir: {tmp}")

    probe_structural_insert_paragraph(tmp)
    probe_structural_remove_paragraph(tmp)
    probe_structural_set_cell_text(tmp)
    probe_structural_insert_table_row(tmp)
    probe_style_edit_bold(tmp)
    probe_mcp_boot_integrity()

    print("\n" + "=" * 60)
    p = sum(1 for s, *_ in results if s == "PASS")
    w = sum(1 for s, *_ in results if s == "WARN")
    f = sum(1 for s, *_ in results if s == "FAIL")
    print(f"PASS={p}  WARN={w}  FAIL={f}")
    return 1 if f else 0


if __name__ == "__main__":
    sys.exit(main())
