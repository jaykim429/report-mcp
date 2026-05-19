"""Round-four probes: mixed edits, MCP transport, pagination edges,
PDF on describe_template, defensive types, large batches."""

from __future__ import annotations

import asyncio
import json
import shutil
import subprocess
import sys
import tempfile
import time
from pathlib import Path

from docx import Document

from report_mcp.server import (
    describe_template,
    fill_and_save,
    inspect_template,
    list_template_targets,
)

P = Path(__file__).parent
PROJECT_ROOT = P.parent
HWPX = next((PROJECT_ROOT / "output" / "templates").glob("*.hwpx"))

findings: list[tuple[str, str, str]] = []


def rec(s: str, pid: str, note: str = "") -> None:
    findings.append((s, pid, note))
    print(f"  {s:4s} {pid}: {note}")


# ────────────────────────────────────────────────────────────────────────
# A. Mixed edits (text + structural + style) in one batch
# ────────────────────────────────────────────────────────────────────────

def probe_mixed_edits(workdir: Path) -> None:
    print("\n[A] mixed TextEdit + StructuralEdit + StyleEdit in one call")
    src = workdir / "mixed.docx"
    out = workdir / "mixed_out.docx"
    d = Document()
    d.add_heading("Status Update", 1)
    d.add_paragraph("[INTRO]")
    d.add_paragraph("Footer line")
    d.save(src)

    paras = list_template_targets(str(src), target_kinds=["paragraph"], max_targets=50).get("targets", [])
    runs = list_template_targets(str(src), target_kinds=["run"], max_targets=50).get("targets", [])
    intro = next(t for t in paras if (t.get("current_text") or "") == "[INTRO]")
    footer = next(t for t in paras if (t.get("current_text") or "") == "Footer line")
    # pick a run inside intro for style edit
    intro_run = next((r for r in runs if (r.get("current_text") or "") == "[INTRO]"), None)
    if intro_run is None:
        rec("WARN", "mixed.norun", "no run target found for [INTRO]")
        return

    edits = [
        {"edit_type": "text", "target_kind": "paragraph",
         "target_id": intro["target_id"], "expected_text_hash": intro["text_hash"],
         "new_text": "이번 주 진행 사항 요약입니다"},
        {"edit_type": "structural", "operation": "insert_paragraph",
         "target_id": footer["target_id"], "position": "before",
         "text": "[새 단락 — 핵심 지표]"},
        {"edit_type": "style", "target_kind": "run",
         "target_id": intro_run["target_id"],
         "bold": True, "font_size_pt": 14.0},
    ]
    r = fill_and_save(str(src), edits, str(out))
    print(f"  status: {r.get('status')}")
    if r.get("status") != "ok":
        rec("FAIL", "mixed.status", f"{r.get('status')} - {r.get('error') or r.get('recovery_hint')}")
        return

    post = Document(out)
    bodies = [p.text for p in post.paragraphs]
    inserted = "[새 단락 — 핵심 지표]" in bodies
    rewritten = "이번 주 진행 사항 요약입니다" in bodies
    # Style: find a run that matches and check bold
    style_ok = False
    for para in post.paragraphs:
        for run in para.runs:
            if "이번 주 진행 사항 요약" in run.text and run.bold:
                style_ok = True
                break
    rec("PASS" if (inserted and rewritten and style_ok) else "FAIL",
        "mixed.combined",
        f"inserted={inserted}, rewritten={rewritten}, style_ok={style_ok}")


# ────────────────────────────────────────────────────────────────────────
# B. Pagination edges
# ────────────────────────────────────────────────────────────────────────

def probe_pagination_edges() -> None:
    print("\n[B] pagination edges (start beyond end, limit=0, start=neg)")
    # total paragraph targets in our HWPX is 87
    total = list_template_targets(str(HWPX), target_kinds=["paragraph"], limit=400)
    n = len(total.get("targets", []))
    print(f"  total paragraph targets: {n}")

    # start beyond end
    r1 = list_template_targets(str(HWPX), target_kinds=["paragraph"], start=n + 10, limit=10)
    print(f"  start={n+10} returned={r1.get('returned')} next_start={r1.get('next_start')}")
    if r1.get("returned") == 0 and r1.get("next_start") is None:
        rec("PASS", "page.beyond", "start beyond end returns empty page, next_start=None")
    else:
        rec("FAIL", "page.beyond", f"unexpected: returned={r1.get('returned')} next_start={r1.get('next_start')}")

    # limit=0 — should return empty page with next_start=0
    r2 = list_template_targets(str(HWPX), target_kinds=["paragraph"], start=0, limit=0)
    print(f"  limit=0 returned={r2.get('returned')} next_start={r2.get('next_start')}")
    if r2.get("returned") == 0:
        rec("PASS", "page.zero_limit", "limit=0 returns empty page")
    else:
        rec("FAIL", "page.zero_limit", f"unexpected returned={r2.get('returned')}")

    # negative start - should be rejected with bad_argument
    r3 = list_template_targets(str(HWPX), target_kinds=["paragraph"], start=-5, limit=5)
    if r3.get("status") == "bad_argument":
        rec("PASS", "page.neg_start", "negative start cleanly rejected with bad_argument")
    else:
        rec("FAIL", "page.neg_start", f"expected bad_argument, got status={r3.get('status')}")


# ────────────────────────────────────────────────────────────────────────
# C. describe_template on PDF — clean format_requires_java?
# ────────────────────────────────────────────────────────────────────────

def probe_describe_pdf(workdir: Path) -> None:
    print("\n[C] describe_template on PDF (no Java)")
    pdf = workdir / "tiny.pdf"
    # Make a 1-page minimal PDF
    try:
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(str(pdf))
        c.drawString(100, 750, "Hello")
        c.save()
    except ImportError:
        rec("WARN", "pdf.skip", "reportlab not installed")
        return

    r = describe_template(str(pdf))
    print(f"  status: {r.get('status')}")
    if r.get("status") == "format_requires_java":
        rec("PASS", "pdf.classify", "clean format_requires_java, no crash")
    elif r.get("status") == "ok":
        rec("PASS", "pdf.actually_works", "JDK is installed and PDF parsed successfully")
    else:
        rec("FAIL", "pdf.bad", f"unexpected status: {r.get('status')} - {r.get('error')}")


# ────────────────────────────────────────────────────────────────────────
# D. Defensive types on edits arg
# ────────────────────────────────────────────────────────────────────────

def probe_defensive_types(workdir: Path) -> None:
    print("\n[D] defensive type handling")
    src = workdir / "def.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "def_out.hwpx"

    # None as edits → should not crash (treat-as-empty is valid)
    try:
        r = fill_and_save(str(src), None, str(out))  # type: ignore[arg-type]
        if r.get("status") == "ok" and r.get("edits_applied", 0) == 0:
            rec("PASS", "def.none", "None edits treated as empty list (status=ok, 0 applied)")
        elif r.get("status") in {"edit_parse_failed", "bad_argument"}:
            rec("PASS", "def.none", f"None edits rejected cleanly: status={r.get('status')}")
        else:
            rec("FAIL", "def.none", f"unexpected: status={r.get('status')} applied={r.get('edits_applied')}")
    except Exception as exc:
        rec("FAIL", "def.none_crash", f"crash on None edits: {type(exc).__name__}: {exc}")

    # edits as wrong type (string)
    try:
        r2 = fill_and_save(str(src), "not a list", str(out))  # type: ignore[arg-type]
        if r2.get("status") != "ok":
            rec("PASS", "def.str", f"non-list edits handled: status={r2.get('status')}")
        else:
            rec("FAIL", "def.str_ok", "string edits returned ok — should reject")
    except Exception as exc:
        rec("FAIL", "def.str_crash", f"crash on string edits: {type(exc).__name__}")

    # edit entry is None
    try:
        r3 = fill_and_save(str(src), [None], str(out))  # type: ignore[list-item]
        if r3.get("status") != "ok":
            rec("PASS", "def.entry_none", f"None inside edits list handled: status={r3.get('status')}")
        else:
            rec("WARN", "def.entry_none_ok", "ok with None in list — odd")
    except Exception as exc:
        rec("FAIL", "def.entry_none_crash", f"crash: {type(exc).__name__}")


# ────────────────────────────────────────────────────────────────────────
# E. Large batch (100 edits) — verify patched lib really scales
# ────────────────────────────────────────────────────────────────────────

def probe_large_batch(workdir: Path) -> None:
    """Stress test: send as many non-conflicting paragraph edits as the
    template has, in one apply call. The patched library should handle
    this without the cumulative _edited FileNotFoundError that previously
    capped us at ~28."""
    print("\n[E] large batch single call (no conflicts)")
    src = workdir / "big.hwpx"
    shutil.copy2(HWPX, src)
    out = workdir / "big_out.hwpx"

    paras = list_template_targets(str(src), target_kinds=["paragraph"], limit=400).get("targets", [])
    edits = [
        {"edit_type": "text", "target_kind": "paragraph",
         "target_id": t["target_id"], "expected_text_hash": t["text_hash"],
         "new_text": f"치환 단락 {i:02d}"}
        for i, t in enumerate(paras)
        if (t.get("current_text") or "").strip()
    ]
    print(f"  prepared {len(edits)} paragraph edits (well past the old 28-edit cap)")
    t0 = time.perf_counter()
    r = fill_and_save(str(src), edits, str(out))
    elapsed = time.perf_counter() - t0
    print(f"  status: {r.get('status')}  edits_applied: {r.get('edits_applied')}  elapsed: {elapsed:.2f}s")
    if r.get("status") == "ok" and r.get("edits_applied") == len(edits):
        rec("PASS", "large.ok", f"{len(edits)} edits in {elapsed:.2f}s, all applied in one call")
    else:
        rec("FAIL", "large.bad",
            f"status={r.get('status')} applied={r.get('edits_applied')} expected={len(edits)}")


# ────────────────────────────────────────────────────────────────────────
# F. MCP transport — actual stdio protocol smoke
# ────────────────────────────────────────────────────────────────────────

def probe_mcp_stdio_transport() -> None:
    """Spawn the actual MCP server, send a `initialize` + `tools/list`
    request over stdio, and verify a well-formed response.

    Uses a background reader thread + queue because Windows can't `select`
    on pipes (only on sockets)."""
    print("\n[F] MCP stdio transport (real protocol roundtrip)")
    py = Path(sys.executable)
    env_extra = {"PYTHONUNBUFFERED": "1", "PYTHONIOENCODING": "utf-8"}

    proc = subprocess.Popen(
        [str(py), "-m", "report_mcp"],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        env={**__import__("os").environ, **env_extra},
        bufsize=0,
    )

    import queue
    import threading

    out_q: "queue.Queue[bytes]" = queue.Queue()

    def reader():
        while True:
            line = proc.stdout.readline()
            if not line:
                return
            out_q.put(line)

    t = threading.Thread(target=reader, daemon=True)
    t.start()

    def send(obj):
        line = (json.dumps(obj) + "\n").encode("utf-8")
        proc.stdin.write(line)
        proc.stdin.flush()

    def recv(timeout: float = 5.0):
        try:
            line = out_q.get(timeout=timeout)
        except queue.Empty:
            return None
        return json.loads(line.decode("utf-8"))

    try:
        # MCP initialize
        send({
            "jsonrpc": "2.0", "id": 1, "method": "initialize",
            "params": {
                "protocolVersion": "2024-11-05",
                "capabilities": {},
                "clientInfo": {"name": "probe-r4", "version": "0.0.1"},
            },
        })
        init_resp = recv()
        if init_resp is None:
            rec("FAIL", "mcp.init", "no response to initialize within timeout")
            return
        if init_resp.get("error"):
            rec("FAIL", "mcp.init", f"initialize error: {init_resp['error']}")
            return
        srv_info = init_resp.get("result", {}).get("serverInfo", {})
        instr_visible = "instructions" in init_resp.get("result", {})
        rec("PASS", "mcp.init", f"server={srv_info.get('name')!r} v={srv_info.get('version')} instructions_in_initialize={instr_visible}")

        # Send `initialized` notification
        send({"jsonrpc": "2.0", "method": "notifications/initialized"})
        time.sleep(0.05)

        # tools/list
        send({"jsonrpc": "2.0", "id": 2, "method": "tools/list"})
        tools_resp = recv()
        if tools_resp is None or tools_resp.get("error"):
            rec("FAIL", "mcp.tools_list", f"tools/list failed: {tools_resp}")
            return
        names = [t["name"] for t in tools_resp["result"]["tools"]]
        expected = {"inspect_template", "list_template_targets", "describe_template", "fill_and_save"}
        missing = expected - set(names)
        if missing:
            rec("FAIL", "mcp.tools_list", f"missing: {missing}, got: {names}")
        else:
            rec("PASS", "mcp.tools_list", f"all 4 tools advertised over JSON-RPC")

        # tools/call — describe_template on our HWPX
        send({
            "jsonrpc": "2.0", "id": 3, "method": "tools/call",
            "params": {"name": "describe_template",
                       "arguments": {"template_path": str(HWPX)}},
        })
        call_resp = recv(timeout=15.0)
        if call_resp is None:
            rec("FAIL", "mcp.tools_call", "no response to tools/call within timeout")
        elif call_resp.get("error"):
            rec("FAIL", "mcp.tools_call", f"error: {call_resp['error']}")
        else:
            # tools/call returns content array; first item is structured JSON for our tool
            content = call_resp["result"].get("content", [])
            structured = call_resp["result"].get("structuredContent")
            payload = structured or (json.loads(content[0]["text"]) if content and content[0].get("type") == "text" else None)
            if payload and payload.get("status") == "ok":
                rec("PASS", "mcp.tools_call",
                    f"describe_template via JSON-RPC: format={payload.get('source_doc_type')} "
                    f"paragraphs={payload.get('total_paragraphs')} tables={payload.get('has_tables')}")
            else:
                rec("FAIL", "mcp.tools_call_payload", f"payload: {payload}")
    finally:
        try:
            proc.stdin.close()
        except Exception:
            pass
        proc.terminate()
        try:
            proc.wait(timeout=3)
        except subprocess.TimeoutExpired:
            proc.kill()


# ────────────────────────────────────────────────────────────────────────

def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="rmcp_r4_"))
    print(f"workdir: {tmp}")

    probe_mixed_edits(tmp)
    probe_pagination_edges()
    probe_describe_pdf(tmp)
    probe_defensive_types(tmp)
    probe_large_batch(tmp)
    probe_mcp_stdio_transport()

    print("\n" + "─" * 70)
    p = sum(1 for s, *_ in findings if s == "PASS")
    w = sum(1 for s, *_ in findings if s == "WARN")
    f = sum(1 for s, *_ in findings if s == "FAIL")
    print(f"PASS={p}  WARN={w}  FAIL={f}")
    for s, pid, note in findings:
        if s != "PASS":
            print(f"  {s} {pid}: {note}")
    return 1 if f else 0


if __name__ == "__main__":
    sys.exit(main())
